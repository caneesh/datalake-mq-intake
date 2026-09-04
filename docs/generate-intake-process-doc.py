"""Generates docs/Message-Intake-Process.docx.

The document is committed as a binary, so this script is committed with it:
edit HERE and regenerate, never edit the .docx by hand. A hand-edited binary
cannot be diffed, cannot be reviewed, and drifts away from its source the first
time someone opens it in a word processor.

    pip install python-docx
    python3 docs/generate-intake-process-doc.py

Deliberately free of any source-system, queue, field or organisation names:
everything is described by the role it plays ("the source queue", "a tracked
binding"), so the document can be handed to anyone. "Storage" means any
distributed file system offering an atomic rename and an explicit durability
barrier -- the two properties the delivery guarantee depends on.

There is a check for that genericness; see the commit that added this file.
Re-run it after any edit:

    python3 - <<'EOF'
    from docx import Document; import re
    d = Document("docs/Message-Intake-Process.docx")
    text = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                text += "\n" + c.text
    banned = ["hdfs", "ibm", "hadoop", "jms", "kerberos", "sequencefile",
              "websphere", "mdb", "spring", "java"]
    print([b for b in banned if re.search(r"\b" + b + r"\b", text, re.I)] or "CLEAN")
    EOF
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

ACCENT = RGBColor(0x0F, 0x5A, 0x64)
MONO = "Consolas"

doc = Document()

# ---------------------------------------------------------------- styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.12

for level, size in ((1, 17), (2, 13), (3, 11.5)):
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = ACCENT
    st.font.bold = True
    st.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    st.paragraph_format.space_after = Pt(5)


def para(text="", style=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def rich(parts, style=None):
    """parts: list of (text, kind) where kind is None|'b'|'i'|'c' (code)."""
    p = doc.add_paragraph(style=style)
    for text, kind in parts:
        run = p.add_run(text)
        if kind == "b":
            run.bold = True
        elif kind == "i":
            run.italic = True
        elif kind == "c":
            run.font.name = MONO
            run.font.size = Pt(9.5)
    return p


def bullets(items):
    for item in items:
        if isinstance(item, list):
            rich(item, style="List Bullet")
        else:
            para(item, style="List Bullet")


def numbered(items):
    for item in items:
        if isinstance(item, list):
            rich(item, style="List Number")
        else:
            para(item, style="List Number")


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            for chunk, kind in (value if isinstance(value, list) else [(value, None)]):
                run = p.add_run(chunk)
                run.font.size = Pt(9.5)
                if kind == "b":
                    run.bold = True
                elif kind == "c":
                    run.font.name = MONO
                    run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def note(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    lab = p.add_run(label + "  ")
    lab.bold = True
    lab.font.color.rgb = ACCENT
    lab.font.size = Pt(10)
    body = p.add_run(text)
    body.font.size = Pt(10)
    body.italic = True
    return p


# ----------------------------------------------------------------- title ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("Message Queue to Distributed Storage Intake")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = ACCENT

sub = doc.add_paragraph()
run = sub.add_run("How messages are consumed, landed and audited")
run.font.size = Pt(13)
run.italic = True

para()
para(
    "This document describes the design of a service that reads messages from a "
    "message queue and writes them to a distributed file system with an "
    "at-least-once delivery guarantee. It explains the consumption cycle, the "
    "ordering rules that make the guarantee hold, how failures are handled, and "
    "how the result is audited and independently checked afterwards."
)
para(
    "It is written generically. Queues, fields and feeds are described by the role "
    "they play rather than by name, so the same description applies to any "
    "deployment of this design. \u201cStorage\u201d throughout means a distributed "
    "file system that supports an atomic rename and an explicit durability barrier "
    "\u2014 the two properties the delivery guarantee depends on. Any store "
    "offering both can host this design."
)

# --------------------------------------------------------------- 1 scope ----
doc.add_heading("1. What the service does", level=1)

para(
    "The service moves messages from a queue onto durable storage without losing "
    "any of them. It does three things, in a fixed order, and the order is the "
    "design:"
)
numbered([
    "Consume a batch of messages from a source queue inside a transaction.",
    "Write those messages to storage and make the file durable and visible.",
    "Acknowledge the messages to the queue manager, and — where configured — "
    "notify a downstream tracker queue in the same transaction.",
])
para(
    "Nothing is acknowledged before the data is safely on storage. If any step "
    "fails before the acknowledgement, the whole batch is rolled back and the "
    "queue manager redelivers it. That is what makes delivery at-least-once."
)

doc.add_heading("1.1 Vocabulary", level=2)
table(
    ["Term", "Meaning"],
    [
        [[("Binding", "b")],
         "One source queue paired with one storage destination, plus the settings that "
         "govern it. A service instance may run several, independently."],
        [[("Listener thread", "b")],
         "A thread that runs the consumption cycle. A binding may have several; each "
         "is fully independent."],
        [[("Batch", "b")],
         "The set of messages consumed and committed as one unit of work. One batch "
         "produces one file."],
        [[("Partition window", "b")],
         "The time bucket a file is filed under on storage. A quarter hour in the "
         "reference configuration."],
        [[("Tracked binding", "b")],
         "A binding that also puts an acknowledgement message on a tracker queue for "
         "each message consumed."],
        [[("Land-only binding", "b")],
         "A binding that writes to storage and nothing else."],
    ],
    widths=[1.5, 4.7],
)

# ------------------------------------------------------------ 2 the rules ---
doc.add_heading("2. The three rules the guarantee rests on", level=1)
para(
    "Everything else in the design is negotiable. These are not, and each exists "
    "because the obvious alternative breaks the guarantee."
)

doc.add_heading("2.1 The message loop is written by hand", level=2)
para(
    "Message-driven containers and template helpers supplied by application "
    "frameworks impose a transaction boundary around every individual message. "
    "That makes batching impossible and multiplies the cost of durability by the "
    "message count. The loop here is written directly against the messaging API so "
    "the transaction boundary is drawn around a whole batch. The application "
    "framework is used for configuration, lifecycle, health and wiring only, and "
    "stays out of the loop entirely."
)

doc.add_heading("2.2 One transacted session per listener thread", level=2)
para(
    "Each listener thread creates its own transacted session, and its consumer and "
    "producers are created from that same session. Sessions are never shared "
    "between threads."
)
para(
    "This is what allows the consume, the tracker put and the poison-message "
    "diversion to be one atomic unit: they are operations on a single session, so "
    "they commit or roll back together. It also means the account the service "
    "connects with must hold read authority on the source queue and write authority "
    "on the tracker and diversion queues, because messaging authenticates at the "
    "connection, not per destination."
)

doc.add_heading("2.3 Acknowledge only after the data is visible", level=2)
para("The write sequence is fixed and never reordered:")
rich([("close", "c"), ("  (the data is durable)  →  ", None),
      ("rename", "c"), ("  (the data is visible)  →  ", None),
      ("commit", "c"), ("  (the messages are acknowledged)", None)])
para(
    "A file is written to a staging directory, forced to disk, closed, and only "
    "then renamed into its final location. Nothing is acknowledged to the queue "
    "manager until that rename has succeeded. A failure at any earlier point rolls "
    "the batch back, and the messages are still on the queue."
)

# --------------------------------------------------------- 3 the cycle ------
doc.add_heading("3. The consumption cycle", level=1)
para(
    "Each listener thread repeats the following loop. Steps 4 to 9 are the unit of "
    "work: everything before the commit can be undone, and nothing after it can."
)

table(
    ["#", "Step", "What it does and why"],
    [
        ["1", "Receive with a timeout",
         "A bounded receive rather than a blocking one, so shutdown can interrupt the "
         "loop predictably."],
        ["2", "Accumulate",
         "The message joins the in-flight batch. Nothing is written yet."],
        ["3", "Test the flush triggers",
         "If any trigger has fired, the batch closes and processing begins. Otherwise "
         "the loop returns to step 1."],
        ["4", "Screen for poison messages",
         "Messages that have been redelivered too many times are diverted to a "
         "separate queue on the same session. A failure to divert rolls the batch back "
         "rather than dropping the message."],
        ["5", "Write to storage",
         "Serialise each message, write to a staging path, force to disk, close, then "
         "rename into the partition."],
        ["6", "Check the balance",
         "Every message consumed must be accounted for as either written or diverted. "
         "An unbalanced batch is rolled back, never committed."],
        ["7", "Send acknowledgements",
         "For a tracked binding, one acknowledgement per message onto the tracker "
         "queue, on the same session."],
        ["8", "Write the audit record",
         "Before the commit, deliberately. See section 6."],
        ["9", "Commit",
         "The single point at which the messages are acknowledged. Everything above "
         "is now irreversible."],
    ],
    widths=[0.35, 1.65, 4.2],
)

note("Why the audit comes before the commit:",
     "written afterwards, a crash in between would leave committed data with no "
     "record, which a balancing control reads as loss. Written first, the same "
     "crash yields an audited file whose messages are redelivered — a duplicate, "
     "which is detectable and true.")

# ------------------------------------------------------- 4 batching ---------
doc.add_heading("4. Batching and when a batch closes", level=1)
para(
    "A batch closes when the first of four triggers fires. All four are always "
    "active except the interval, which can be disabled."
)
table(
    ["Trigger", "Fires when", "Purpose"],
    [
        ["Count", "The batch reaches the configured message count.",
         "Bounds the size of a unit of work."],
        ["Size", "The accumulated payload reaches the configured byte budget.",
         "Bounds memory, and keeps output files to a workable size."],
        ["Partition boundary", "The clock leaves the window the batch opened in.",
         "Keeps a batch to a single time bucket, and keeps a quiet period to one file "
         "rather than one file per message."],
        ["Interval", "A configured time has passed since the batch's first message.",
         "Bounds how long a message waits. Optional."],
    ],
    widths=[1.2, 2.4, 2.6],
)

para(
    "The partition trigger cannot be switched off. Without it, a low-volume feed "
    "produces one small file per interval, and at trickle volume approaches one "
    "file per message — a pattern that has been observed to overwhelm the storage "
    "layer and the messaging listeners together. Bounding a batch to one window "
    "puts a floor under the file cadence regardless of how the other triggers are "
    "tuned."
)
para(
    "The interval is measured from the batch's first message, not from when the "
    "previous batch was flushed. Measured from the flush, an idle period longer "
    "than the interval leaves the next message already expired, so it is written "
    "alone — reproducing exactly the one-file-per-message pattern the trigger "
    "exists to prevent."
)

# ------------------------------------------------------- 5 writing ----------
doc.add_heading("5. Writing to storage", level=1)

doc.add_heading("5.1 File format", level=2)
para(
    "Records are written to a container file format that holds a sequence of "
    "key/value pairs. Both the key type and the value type form a contract with "
    "downstream consumers: the file header embeds the type names, so changing "
    "either produces files that existing readers cannot open. The declared types "
    "are pinned by an automated check for that reason."
)
para(
    "Record-level or no compression is used. Block compression is deliberately "
    "avoided: it triggers mid-stream buffering behaviour that interacts badly with "
    "erasure-coded storage."
)

doc.add_heading("5.2 Staging, durability and visibility", level=2)
numbered([
    "The file is created under a staging directory that is private to the process "
    "instance, so two instances on one host can never write to the same path.",
    "Records are appended one at a time.",
    "The stream is forced to disk once per batch, not once per record. This is the "
    "difference between surviving a process crash and surviving a correlated power "
    "loss after the messages have already been acknowledged.",
    "The file is closed.",
    "The file is renamed into its partition directory. The rename is what makes it "
    "visible; a reader scanning the partition never sees a partial file.",
])
para(
    "Any failure before the rename deletes the staged file, so a failed batch "
    "leaves nothing behind. Staged files older than a configured age are also swept "
    "at startup, which cleans up after a crash."
)

doc.add_heading("5.3 Partition layout", level=2)
para(
    "Files are filed under a directory path derived from the clock — year, month, "
    "day, hour and sub-hour window. The path is computed fresh at every flush and "
    "never cached; a cached path was the confirmed cause of a production incident "
    "in which files landed in a bucket well over an hour stale."
)
note("A contract decision worth knowing:",
     "the partition is chosen from the clock at the moment of the write. Because "
     "the partition trigger fires just after a window closes, a batch that "
     "accumulated in one window is filed under the next. The alternative — filing "
     "under the window the messages arrived in — writes into a bucket that has "
     "just closed, which a downstream job that sweeps each bucket once would miss "
     "entirely. Filing forward keeps every write landing in a bucket that is still "
     "open. Confirm which behaviour downstream consumers expect.")

para(
    "File names combine the binding, the process instance, a timestamp and a "
    "per-instance sequence number, so two instances writing concurrently can never "
    "collide."
)

# ------------------------------------------------------ 6 failures ----------
doc.add_heading("6. What happens when something fails", level=1)

doc.add_heading("6.1 Before the commit", level=2)
para(
    "Any failure before the commit rolls the transaction back. Every message in the "
    "batch returns to the queue and is redelivered. Nothing is lost, and no partial "
    "batch is ever committed. If a file had already been renamed into place, the "
    "redelivered messages produce a second copy — a duplicate, which the design "
    "permits and reconciliation can detect."
)

doc.add_heading("6.2 After the commit", level=2)
para(
    "A failure in bookkeeping after the commit must not roll anything back: the "
    "messages are already acknowledged and can never be redelivered. Such failures "
    "are logged and contained. Treating them as rollbacks would corrupt the "
    "internal state that tracks suspect messages and would permanently reduce "
    "throughput."
)

doc.add_heading("6.3 Classifying failures", level=2)
para(
    "Not every failure means the same thing, and the response differs. Failures are "
    "classified before anything is decided:"
)
table(
    ["Class", "Examples", "Response"],
    [
        ["Message data", "A payload that cannot be serialised or parsed.",
         "Enter degraded mode and begin isolating the offending message."],
        ["Storage infrastructure", "The storage layer is unwritable or unreachable.",
         "Retry. Do not shrink batches; the messages are not at fault."],
        ["Messaging infrastructure", "A broken session or connection; a full queue.",
         "Rebuild the session with bounded backoff, or stall and alert."],
        ["Security or configuration", "Bad credentials, denied authority.",
         "Stop. Retrying cannot fix it, and retrying hides the real error."],
        ["Unknown", "Anything unclassified.",
         "Roll back and alert. Never assume a message is at fault."],
    ],
    widths=[1.4, 2.2, 2.6],
)

doc.add_heading("6.4 Isolating a bad message", level=2)
para(
    "When a failure is classified as message data, the binding enters a degraded "
    "mode in which the batch size is reduced — either to a single message, or by "
    "halving repeatedly. The messages of the failed batch are recorded as suspects "
    "by their message identifiers, shared across all of that binding's listener "
    "threads."
)
para(
    "Tracking by identifier rather than by thread matters: after a rollback the "
    "queue manager is free to redeliver the batch's messages to any listener, so "
    "an approach that assumed the same thread would see them again would not work. "
    "Clean subsets commit normally and clear themselves from the suspect set; "
    "batches containing the genuine offender keep failing and shrinking until it is "
    "alone in its own unit of work. Normal batch size is restored only after the "
    "suspect set is empty and a configured number of consecutive successes."
)

doc.add_heading("6.5 Poison messages", level=2)
para(
    "The queue manager increments a delivery count each time a message is rolled "
    "back. When that count passes a configured threshold, the message is diverted "
    "to a separate queue on the same transacted session, so the diversion and the "
    "acknowledgement are atomic. A failure to divert rolls the batch back rather "
    "than dropping the message."
)
note("An important subtlety:",
     "delivery count alone cannot distinguish a genuinely bad message from a "
     "healthy one that happened to sit in several batches which rolled back for "
     "an unrelated reason. Diversion can therefore be gated so it only applies "
     "while failures look like message data. Without that gate, a storage outage "
     "lasting a few retry cycles diverts whole batches of healthy messages, which "
     "then need manual replay.")

doc.add_heading("6.6 Session recovery", level=2)
para(
    "A broken session is closed and rebuilt against the same connection, with "
    "exponential backoff and jitter, up to a bounded number of attempts. Failures "
    "that recovery cannot fix — bad credentials, denied authority, an unknown "
    "destination — are identified and stop the loop immediately rather than "
    "consuming the retry budget."
)
para(
    "Identifying them requires care: messaging clients commonly report every "
    "connection failure with the same top-level error, and put the reason that "
    "distinguishes them in a nested exception. Matching only the top-level message "
    "silently recognises none of them."
)

# --------------------------------------------------------- 7 auditing ------
doc.add_heading("7. Auditing", level=1)
para(
    "The audit trail exists so that a question like 'did everything that was "
    "consumed actually land?' can be answered from records, not from logs."
)

doc.add_heading("7.1 The audit record", level=2)
para(
    "One record is written per unit of work, before the commit, to its own "
    "immutable file. Each record names the file it accounts for and carries:"
)
bullets([
    "the binding and the partition the data was written to",
    "the number of records written and the number of bytes",
    "the number of messages consumed from the queue",
    "the number of messages diverted as poison",
    "the difference between those figures, and whether it balances",
    "the identity of the first and last record, where identity is available",
    "the process instance and the commit timestamp",
])
para(
    "Records are staged and renamed like data files, so a control can never read a "
    "half-written record. Whether a failure to write the audit stops the batch is "
    "configurable and defaults to stopping it: committing data that no record "
    "accounts for produces exactly the gap the audit exists to close."
)

doc.add_heading("7.2 The transaction-time balance check", level=2)
para(
    "Before every commit, three independently observed numbers must agree:"
)
rich([("messages consumed", "b"), ("  =  ", None), ("records written", "b"),
      ("  +  ", None), ("messages diverted", "b")])
para(
    "Each figure comes from a different place: the count taken from the queue, the "
    "count the writer observed itself append, and the count the poison screen "
    "observed itself divert. None is computed from another. An unbalanced batch is "
    "rolled back and redelivered rather than committed, so a silently dropped "
    "message becomes a redelivery instead of a loss."
)
note("Why independence matters:",
     "a check that compares a number with itself always passes. If two of these "
     "figures were derived from the same source, the check would be decorative. "
     "Keeping them independent observations is the whole value of the control.")

doc.add_heading("7.3 The record index", level=2)
para(
    "Where a payload carries a stable per-message identifier, a small index file is "
    "written alongside the data file listing each record's offset and identity. It "
    "is written after the data file is visible and never fails the batch: refusing "
    "to commit because a reconciliation aid could not be written would roll back "
    "landed data and manufacture a duplicate."
)
para(
    "The index is optional per binding, and should stay off where no trustworthy "
    "identifier exists. An index of blanks is worse than no index, because the "
    "unidentified records read as losses."
)

# ------------------------------------------------- 8 reconciliation ---------
doc.add_heading("8. Reconciliation", level=1)
para(
    "The audit says what should be on storage. Reconciliation confirms it actually "
    "is. It runs on a schedule, reads only, holds no messaging session, and cannot "
    "affect consumption in any way — a checking mechanism that can halt the thing "
    "it checks is worse than no check at all."
)

doc.add_heading("8.1 What a pass does", level=2)
numbered([
    "Wait until a partition has closed and a grace period has passed, so batches "
    "still in flight are not mistaken for missing data.",
    "List the files actually present in the partition.",
    "Read the audit records that claim to describe that partition.",
    "Compare them.",
])

table(
    ["Finding", "Meaning"],
    [
        ["Audit record with no file",
         "The loudest result: something was recorded as landed and is not there."],
        ["Count mismatch",
         "The file exists but holds a different number of records than the audit "
         "claims — truncation or damage after landing."],
        ["Unreadable file",
         "The file cannot be opened. Reported and retried, never ignored."],
        ["Unreadable audit record",
         "The scan itself was incomplete, so no clean verdict is possible."],
        ["File with no audit record",
         "Classified by comparing its record identities against other files: either a "
         "duplicate of data already accounted for, or the only copy."],
    ],
    widths=[1.7, 4.5],
)

doc.add_heading("8.2 Rules it follows", level=2)
bullets([
    "The record count is always read from the data file itself, never taken from "
    "the index. The index and the audit record are produced together by the same "
    "write, so trusting the index would mean comparing a number with itself.",
    "Nothing is ever deleted. The only corrective action available is moving a "
    "file classified as a duplicate into a quarantine area, and even that is off by "
    "default — a file is reported, not moved, without a human decision.",
    "A binding with no approved record identity refuses to reconcile rather than "
    "guessing. A confidently wrong answer is worse than no answer.",
    "A partition that could not be resolved is carried forward and re-examined on "
    "later passes, including across restarts, rather than dropping off the schedule "
    "once it ages out of the recent window.",
    "Each binding is reconciled independently and on its own worker, so one slow or "
    "failing binding cannot delay another.",
])

# ------------------------------------------------- 9 observability ----------
doc.add_heading("9. Health and observability", level=1)

doc.add_heading("9.1 Health semantics", level=2)
para(
    "This service accepts no inbound traffic, so reporting itself unhealthy "
    "achieves nothing except a restart. Health status is therefore graded, and a "
    "hard failure status is reserved for the case where every binding is down:"
)
table(
    ["Status", "Meaning", "Effect"],
    [
        ["Healthy", "All bindings consuming normally.", "None."],
        ["Degraded", "A binding is isolating a bad message, recovering a session, or "
                     "rolling back repeatedly.", "Alert. No restart."],
        ["Partial outage", "Some bindings are down while others still work.",
         "Alert. Restarting would interrupt the healthy ones to fix one that is "
         "already contained."],
        ["Down", "Every binding is unhealthy.", "The process genuinely has nothing to do."],
    ],
    widths=[1.1, 3.0, 2.1],
)

doc.add_heading("9.2 Signals worth alerting on", level=2)
table(
    ["Signal", "Condition", "What it means"],
    [
        ["Balance check failures", "Any increase at all",
         "A message was dropped before a commit. The most serious signal available."],
        ["Diversion queue depth", "Non-zero",
         "Messages are sitting undelivered and need manual attention."],
        ["Acknowledgements sent", "Flat while consumption rises",
         "Data is landing without being acknowledged downstream."],
        ["Binding health", "Not healthy for a sustained period",
         "A binding is stalled. Nothing is lost, but nothing is progressing either."],
        ["Reconciliation discrepancies", "Any increase",
         "Landed data and the audit trail disagree."],
        ["Identity extraction misses", "Sustained increase",
         "Payloads are arriving without their identifier — an upstream change."],
    ],
    widths=[1.5, 1.8, 2.9],
)
para(
    "A positive signal is as important as a failure signal. A counter that only "
    "increments on failure cannot distinguish 'working' from 'not running', so the "
    "count of successful acknowledgements is published for exactly that reason."
)

doc.add_heading("9.3 Listener supervision", level=2)
para(
    "Each listener thread is watched. A thread that dies unexpectedly is detected "
    "and reported rather than silently reducing throughput while the service "
    "continues to report itself healthy."
)

# ----------------------------------------------------- 10 startup -----------
doc.add_heading("10. Startup checks", level=1)
para(
    "The service refuses to start in states that would otherwise fail silently at "
    "runtime, or worse, appear to succeed:"
)
bullets([
    "Configuration is validated as a whole — queue names must not collide, batch "
    "sizes must fit inside the queue manager's uncommitted-message limit, and the "
    "combined memory budget of all bindings must fit in the heap.",
    "Storage destinations, including the audit destination, are proved writable. "
    "Otherwise the service starts and then stalls on its first batch.",
    "The storage cluster identity is checked against an expected value. Pointed at "
    "the wrong cluster, everything else succeeds — it connects, authenticates, "
    "writes, and reports healthy, against the wrong destination.",
    "Placeholder or provisional components are refused in a production posture "
    "unless explicitly accepted for a named binding.",
    "Staged files left by a previous run of the same instance are swept.",
])
para(
    "A separate pre-flight mode connects to every dependency using the deployment's "
    "own configuration, checks one fact at a time, prints a report naming the remedy "
    "for each failure, and exits. It consumes nothing, sends nothing to any queue a "
    "downstream system reads, and writes only inside its own staging area, so it is "
    "safe to run against an environment carrying live data."
)

# ------------------------------------------------ 11 guarantees -------------
doc.add_heading("11. Guarantees, and what they are not", level=1)
table(
    ["Property", "Position"],
    [
        [[("Delivered at least once", "b")],
         "Guaranteed. No acknowledged message is ever lost."],
        [[("Delivered exactly once", "b")],
         "Not provided. A failure after a file is visible but before the commit "
         "produces a duplicate. This is deliberate: the alternative is risking loss."],
        [[("Duplicates detectable", "b")],
         "Where a stable per-message identity exists, yes, via the audit trail and "
         "reconciliation. Without one, duplicates can be reported but not classified."],
        [[("Order preserved", "b")],
         "Not across listener threads. Several threads consume concurrently, so the "
         "order records appear in files is not the order they were produced."],
        [[("No partial batches", "b")],
         "Guaranteed. A unit of work commits entirely or not at all."],
        [[("No partial files visible", "b")],
         "Guaranteed. Files become visible by rename, never by being written in place."],
    ],
    widths=[1.7, 4.5],
)

doc.add_heading("11.1 Where duplicates can arise", level=2)
para(
    "There is exactly one window, and it is worth stating plainly. The file is made "
    "visible before the messages are acknowledged. If the process fails between "
    "those two events, the data is on storage and the messages are still on the "
    "queue, so they will be delivered again and written again."
)
para(
    "This is the correct trade. The alternative ordering — acknowledge first, then "
    "write — turns the same crash into permanent loss. A duplicate can be found and "
    "removed; a loss cannot be recovered."
)

# ------------------------------------------------ 12 tuning -----------------
doc.add_heading("12. What can be tuned, and what it affects", level=1)
table(
    ["Setting", "Affects"],
    [
        ["Batch message count", "Unit-of-work size. Must fit inside the queue "
                                "manager's uncommitted-message limit — for a tracked "
                                "binding each message costs two operations, so the "
                                "usable limit is halved."],
        ["Batch byte budget", "File size and memory. The live budget is this figure "
                              "multiplied by the number of listener threads, and "
                              "retained memory runs higher than the payload estimate."],
        ["Batch interval", "How long a message may wait. Disabling it leaves the "
                           "partition boundary as the only time-based trigger."],
        ["Listener threads", "Throughput and concurrency, and a multiplier on the "
                             "memory budget."],
        ["Diversion threshold", "How many redeliveries a message survives before "
                                "being diverted. Must be set consistently on the "
                                "queue itself. Size it above plausible outage "
                                "durations, or an outage diverts healthy messages."],
        ["Degradation strategy", "How aggressively batches shrink while isolating a "
                                 "bad message."],
        ["Balance check", "Whether an unbalanced batch is refused. Recommended on "
                          "for any feed that must not lose data."],
        ["Durability mode", "Whether each batch is forced to disk before close, or "
                            "only flushed to the storage pipeline."],
        ["Record index", "Whether identity metadata is written alongside data files. "
                         "Requires a trustworthy identifier."],
        ["Reconciliation", "Whether landed data is checked against the audit trail, "
                           "and how far back each pass looks."],
    ],
    widths=[1.6, 4.6],
)

# ------------------------------------------------ closing -------------------
doc.add_heading("13. Summary", level=1)
para(
    "The design is built around a single ordering rule — durable, then visible, "
    "then acknowledged — and everything else follows from protecting it. Batching "
    "makes the write efficient without weakening the rule. Classification decides "
    "whether a failure means the data is at fault or the environment is. The audit "
    "trail records what happened at the moment it happened, and reconciliation "
    "checks that record against reality later, independently."
)
para(
    "The result is a service that will duplicate a record under a narrow and "
    "well-understood failure, and will not lose one."
)

out = "/home/aneesh/projects/datalake-mq-intake/docs/Message-Intake-Process.docx"
doc.save(out)
print("written:", out)
